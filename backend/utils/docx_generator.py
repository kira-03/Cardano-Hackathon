"""
Word Document Generator for Exchange Listing Proposals
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
import logging
from typing import Dict, Any, List

logger = logging.getLogger(__name__)


class ListingProposalGenerator:
    """Generate professional Word documents for exchange listing proposals"""
    
    def __init__(self):
        self.output_dir = "outputs/proposals"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_proposal(self, analysis_id: str, proposal_data: Dict[str, Any]) -> str:
        """
        Generate a professional exchange listing proposal as a Word document
        
        Args:
            analysis_id: Unique identifier for this analysis
            proposal_data: Complete proposal data from exchange preparation agent
            
        Returns:
            Path to generated .docx file
        """
        try:
            logger.info(f"📄 Generating exchange listing proposal for {analysis_id}")
            
            doc = Document()
            
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            self._add_title(doc, proposal_data)
            self._add_executive_summary(doc, proposal_data)
            self._add_token_overview(doc, proposal_data)
            self._add_metrics_section(doc, proposal_data)
            self._add_readiness_assessment(doc, proposal_data)
            self._add_compliance_checklist(doc, proposal_data)
            self._add_market_section(doc, proposal_data)
            self._add_contact_section(doc, proposal_data)
            
            filename = f"listing_proposal_{analysis_id[:8]}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            logger.info(f"✅ Proposal generated: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"❌ Error generating proposal: {e}", exc_info=True)
            return None
    
    def _add_title(self, doc: Document, data: Dict[str, Any]):
        """Add document title and header"""
        title = doc.add_heading('Exchange Listing Proposal', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(title, 12, 12)
        
        sTokenName = data.get("token_name", "Unknown")
        sTokenSymbol = data.get("token_symbol", "N/A")
        token_title = doc.add_heading(f'{sTokenName} ({sTokenSymbol})', 1)
        token_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(token_title, 12, 12)
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'Generated: {datetime.utcnow().strftime("%B %d, %Y")}')
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        self._set_spacing(date_para, 6, 18)
    
    def _add_executive_summary(self, doc: Document, data: Dict[str, Any]):
        """Add executive summary section"""
        heading = doc.add_heading('Executive Summary', 1)
        self._set_spacing(heading, 18, 12)
        
        oReadinessScore = data.get("readiness_score", {})
        sGrade = oReadinessScore.get("grade", "N/A")
        nScore = oReadinessScore.get("total", 0)
        nCompliance = data.get("compliance_rate", 0)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run('Listing Readiness Grade: ').bold = True
        grade_run = summary_para.add_run(f'{sGrade} ({nScore:.1f}/100)')
        grade_run.font.size = Pt(14)
        grade_run.font.color.rgb = RGBColor(0, 102, 204)
        grade_run.bold = True
        self._set_spacing(summary_para, 6, 6)
        
        compliance_para = doc.add_paragraph()
        compliance_para.add_run('Exchange Requirements Met: ').bold = True
        compliance_run = compliance_para.add_run(f'{nCompliance:.1f}%')
        compliance_run.font.color.rgb = RGBColor(34, 139, 34) if nCompliance >= 70 else RGBColor(204, 51, 0)
        compliance_run.bold = True
        self._set_spacing(compliance_para, 6, 12)
        
        sUvp = data.get("unique_value_proposition", "")
        if sUvp:
            spacer = doc.add_paragraph()
            self._set_spacing(spacer, 6, 6)
            
            uvp_heading = doc.add_paragraph()
            uvp_heading.add_run('Unique Value Proposition').bold = True
            self._set_spacing(uvp_heading, 6, 6)
            
            quote = doc.add_paragraph(sUvp, style='Intense Quote')
            self._set_spacing(quote, 6, 12)
        
        doc.add_page_break()
    
    def _add_token_overview(self, doc: Document, data: Dict[str, Any]):
        """Add token overview section"""
        heading = doc.add_heading('1. Token Overview', 1)
        self._set_spacing(heading, 18, 12)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = 'Property'
        cells[1].text = 'Value'
        self._set_cell_bold(cells[0])
        self._set_cell_bold(cells[1])
        
        table.rows[1].cells[0].text = 'Token Name'
        table.rows[1].cells[1].text = data.get("token_name", "N/A")
        
        table.rows[2].cells[0].text = 'Symbol'
        table.rows[2].cells[1].text = data.get("token_symbol", "N/A")
        
        table.rows[3].cells[0].text = 'Blockchain'
        table.rows[3].cells[1].text = 'Cardano'
        
        table.rows[4].cells[0].text = 'Policy ID'
        sPolicyText = data.get("policy_id") or "N/A"
        if sPolicyText and sPolicyText != "N/A" and len(sPolicyText) > 50:
            sPolicyText = f"{sPolicyText[:25]}...{sPolicyText[-15:]}"
        table.rows[4].cells[1].text = sPolicyText
        
        spacer = doc.add_paragraph()
        self._set_spacing(spacer, 12, 12)
        
        desc_heading = doc.add_paragraph()
        desc_heading.add_run('Description').bold = True
        self._set_spacing(desc_heading, 6, 6)
        
        sDescription = data.get("description", "No description available")
        desc_para = doc.add_paragraph(sDescription)
        self._set_spacing(desc_para, 6, 12)
        
        oSocial = data.get("social_links", {})
        if any(oSocial.values()):
            spacer = doc.add_paragraph()
            self._set_spacing(spacer, 6, 6)
            
            social_heading = doc.add_paragraph()
            social_heading.add_run('Official Links').bold = True
            self._set_spacing(social_heading, 6, 6)
            
            aLinkList = []
            if data.get("website") and data.get("website") != "N/A":
                aLinkList.append(f'Website: {data.get("website")}')
            if oSocial.get("twitter") and oSocial.get("twitter") != "N/A":
                aLinkList.append(f'Twitter: {oSocial.get("twitter")}')
            if oSocial.get("telegram") and oSocial.get("telegram") != "N/A":
                aLinkList.append(f'Telegram: {oSocial.get("telegram")}')
            if oSocial.get("discord") and oSocial.get("discord") != "N/A":
                aLinkList.append(f'Discord: {oSocial.get("discord")}')
            
            for sLink in aLinkList:
                p = doc.add_paragraph(sLink, style='List Bullet')
                self._set_spacing(p, 3, 3)
        
        doc.add_page_break()
    
    def _add_metrics_section(self, doc: Document, data: Dict[str, Any]):
        """Add key metrics section"""
        heading = doc.add_heading('2. Key Metrics', 1)
        self._set_spacing(heading, 18, 12)
        
        oMetrics = data.get("metrics", {})
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = 'Metric'
        cells[1].text = 'Value'
        self._set_cell_bold(cells[0])
        self._set_cell_bold(cells[1])
        
        table.rows[1].cells[0].text = 'Total Supply'
        table.rows[1].cells[1].text = str(oMetrics.get("total_supply", "N/A"))
        
        table.rows[2].cells[0].text = 'Total Holders'
        nHolders = oMetrics.get("holders", "N/A")
        table.rows[2].cells[1].text = f'{nHolders:,}' if isinstance(nHolders, int) else str(nHolders)
        
        table.rows[3].cells[0].text = 'Liquidity (USD)'
        table.rows[3].cells[1].text = str(oMetrics.get("liquidity", "N/A"))
        
        table.rows[4].cells[0].text = '24h Trading Volume'
        table.rows[4].cells[1].text = str(oMetrics.get("volume_24h", "N/A"))
        
        table.rows[5].cells[0].text = 'Top 10 Holders Concentration'
        table.rows[5].cells[1].text = str(oMetrics.get("top_10_concentration", "N/A"))
        
        doc.add_page_break()
    
    def _add_readiness_assessment(self, doc: Document, data: Dict[str, Any]):
        """Add readiness assessment section"""
        heading = doc.add_heading('3. Listing Readiness Assessment', 1)
        self._set_spacing(heading, 18, 12)
        
        oReadiness = data.get("readiness_score", {})
        oBreakdown = oReadiness.get("breakdown", {})
        
        para = doc.add_paragraph()
        para.add_run('Overall Readiness: ').bold = True
        sGrade = oReadiness.get("grade", "N/A")
        nTotal = oReadiness.get("total", 0)
        score_run = para.add_run(f'{sGrade} Grade ({nTotal:.1f}/100)')
        score_run.font.size = Pt(12)
        score_run.font.color.rgb = RGBColor(0, 102, 204)
        self._set_spacing(para, 6, 12)
        
        if oBreakdown:
            breakdown_heading = doc.add_paragraph('Score Breakdown:', style='Heading 3')
            self._set_spacing(breakdown_heading, 12, 6)
            
            for sCategory, nScore in oBreakdown.items():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f'{sCategory}: ').bold = True
                sScoreText = f'{nScore:.1f}/100'
                score_run = p.add_run(sScoreText)
                
                if nScore >= 80:
                    score_run.font.color.rgb = RGBColor(34, 139, 34)
                elif nScore >= 60:
                    score_run.font.color.rgb = RGBColor(255, 140, 0)
                else:
                    score_run.font.color.rgb = RGBColor(204, 51, 0)
                
                self._set_spacing(p, 3, 3)
        
        doc.add_page_break()
    
    def _add_compliance_checklist(self, doc: Document, data: Dict[str, Any]):
        """Add exchange requirements compliance checklist"""
        heading = doc.add_heading('4. Exchange Requirements Compliance', 1)
        self._set_spacing(heading, 18, 12)
        
        aRequirements = data.get("requirements", [])
        
        if not aRequirements:
            empty = doc.add_paragraph("No specific requirements data available.")
            self._set_spacing(empty, 6, 6)
            return
        
        oExchanges = {}
        for oReq in aRequirements:
            sExchange = oReq.get("exchange", "Unknown")
            if sExchange not in oExchanges:
                oExchanges[sExchange] = []
            oExchanges[sExchange].append(oReq)
        
        for sExchange, aReqs in oExchanges.items():
            exchange_heading = doc.add_paragraph(f'{sExchange} Requirements:', style='Heading 3')
            self._set_spacing(exchange_heading, 12, 6)
            
            for oReq in aReqs:
                p = doc.add_paragraph(style='List Bullet')
                
                sCheckbox = '☑' if oReq.get("meets_requirement") else '☐'
                p.add_run(f'{sCheckbox} ').bold = True
                p.add_run(oReq.get("requirement", "N/A"))
                
                sStatusText = f' ({oReq.get("current_status", "N/A")})'
                status_run = p.add_run(sStatusText)
                status_run.font.size = Pt(9)
                status_run.font.color.rgb = RGBColor(128, 128, 128)
                
                self._set_spacing(p, 3, 3)
            
            spacer = doc.add_paragraph()
            self._set_spacing(spacer, 6, 6)
        
        doc.add_page_break()
    
    def _add_market_section(self, doc: Document, data: Dict[str, Any]):
        """Add market potential section"""
        heading = doc.add_heading('5. Market Potential & Growth Strategy', 1)
        self._set_spacing(heading, 18, 12)
        
        sMarketPotential = data.get("market_potential", "")
        if sMarketPotential:
            market_para = doc.add_paragraph(sMarketPotential)
            self._set_spacing(market_para, 6, 12)
        
        aRecommended = data.get("recommended_exchanges", [])
        if aRecommended:
            spacer = doc.add_paragraph()
            self._set_spacing(spacer, 6, 6)
            
            exchanges_heading = doc.add_paragraph('Recommended Target Exchanges:', style='Heading 3')
            self._set_spacing(exchanges_heading, 12, 6)
            
            for sExchange in aRecommended:
                p = doc.add_paragraph(sExchange, style='List Bullet')
                self._set_spacing(p, 3, 3)
        
        doc.add_page_break()
    
    def _add_contact_section(self, doc: Document, data: Dict[str, Any]):
        """Add contact information section"""
        heading = doc.add_heading('6. Contact Information', 1)
        self._set_spacing(heading, 18, 12)
        
        intro_para = doc.add_paragraph('For additional information or questions regarding this listing proposal, please contact:')
        self._set_spacing(intro_para, 6, 12)
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light List Accent 1'
        
        table.rows[0].cells[0].text = 'Project Website'
        table.rows[0].cells[1].text = data.get("website", "N/A")
        
        oSocial = data.get("social_links", {})
        table.rows[1].cells[0].text = 'Twitter'
        table.rows[1].cells[1].text = oSocial.get("twitter", "N/A")
        
        table.rows[2].cells[0].text = 'Telegram'
        table.rows[2].cells[1].text = oSocial.get("telegram", "N/A")
        
        spacer = doc.add_paragraph()
        self._set_spacing(spacer, 12, 12)
        
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('This document was generated by Cross-Chain Navigator AI')
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        self._set_spacing(footer, 12, 6)
    
    def _set_cell_bold(self, cell):
        """Make cell text bold"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    def _set_spacing(self, paragraph, space_before: int = 0, space_after: int = 0):
        """Set paragraph spacing before and after"""
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
